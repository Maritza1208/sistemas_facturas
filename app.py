from flask import Flask, request, render_template, redirect, url_for, session, jsonify, send_file
import os
import json
import pandas as pd
import base64
import xml.etree.ElementTree as ET
import requests
import io
import zipfile
import shutil
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from collections import defaultdict
from reportlab.lib.units import inch
from reportlab.platypus import Table, TableStyle
from docx.oxml.ns import qn
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from lxml import etree
from docx import Document
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from flask_session import Session
from docx.shared import Inches
import urllib3
import re
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)


# ========== Configuración ==========
UPLOAD_FOLDER = "uploads"
HISTORIAL_PATH = os.path.join(UPLOAD_FOLDER, "corregidas.json")

TEMP_UPLOADS_FOLDER = os.path.join(os.getcwd(), "temp_uploads")
os.makedirs(TEMP_UPLOADS_FOLDER, exist_ok=True)

INCOMING_FOLDER = TEMP_UPLOADS_FOLDER 

if not os.path.exists(HISTORIAL_PATH):
    with open(HISTORIAL_PATH, 'w', encoding='utf-8') as f:
        json.dump({}, f, indent=2, ensure_ascii=False)

API_LOGIN_URL = "https://172.17.100.104:9443/api/Auth/LoginSISPRO"
API_CARGA_JSON_URL = "https://172.17.100.104:9443/api/PaquetesFevRips/CargarFevRips"
API_CREDENCIALES = {
    "persona": {
        "identificacion": {
            "tipo": "CC",
            "numero": "1085301187"
        }
    },
    "clave": "HilaSistemas2024*",
    "nit": "891200240"
}

# ========== Funciones ==========

def limpiar_archivos_sin_cuv(upload_folder):
    archivos = os.listdir(upload_folder)
   
    facturas_con_cuv = set()

    # Encontrar facturas corregidas (tienen archivo _CUV_CORREGIDO.json)
    for archivo in archivos:
        if archivo.endswith("_CUV_CORREGIDO.json"):
            num_factura = archivo.split("_")[0]
            facturas_con_cuv.add(num_factura)

    # Construir lista de archivos a conservar (los de las facturas corregidas)
    archivos_a_conservar = set()
    for factura in facturas_con_cuv:
        for archivo in archivos:
            if archivo.startswith(factura + "_"):
                archivos_a_conservar.add(archivo)

    # Siempre conservar corregidas.json
    archivos_a_conservar.add("corregidas.json")

    # Eliminar todo archivo que NO esté en la lista para conservar
    for archivo in archivos:
        if archivo not in archivos_a_conservar:
            ruta_completa = os.path.join(upload_folder, archivo)
            if os.path.isfile(ruta_completa):
                os.remove(ruta_completa)
                print(f"🗑️ Archivo eliminado: {archivo}")

def limpiar_num_factura(num):
    """Elimina todo lo que no sea número del código de factura."""
    return re.sub(r"[^0-9]", "", num)


def buscar_attdoc(num_factura_limpio, carpeta_uploads):
    for archivo in os.listdir(carpeta_uploads):
        if archivo.endswith("_2_AttDoc.xml") and num_factura_limpio in archivo:
            return os.path.join(carpeta_uploads, archivo)
    return None


def corregir_json_valido(ruta_json_original, ruta_salida, uploads_path, num_factura):
    try:
        with open(ruta_json_original, "r", encoding="utf-8") as f:
            data = json.load(f)

        # Ruta esperada del archivo AttDoc.xml
        ruta_attdoc = os.path.join(uploads_path, f"{num_factura}_2_AttDoc.xml")

        if not os.path.exists(ruta_attdoc):
            print(f"❌ No se encontró el archivo AttDoc: {ruta_attdoc}")
            return False

        # Leer el XML en binario y codificar a Base64
        with open(ruta_attdoc, "rb") as archivo_xml:
            contenido_xml = archivo_xml.read()
            contenido_base64 = base64.b64encode(contenido_xml).decode("utf-8")

        # Reemplazar el campo xmlFevFile
        data["xmlFevFile"] = contenido_base64

        # Guardar el JSON corregido
        with open(ruta_salida, "w", encoding="utf-8") as fout:
            json.dump(data, fout, indent=2, ensure_ascii=False)

        print(f"✅ JSON corregido guardado: {ruta_salida}")
        return True

    except Exception as e:
        print(f"❌ Error al corregir JSON para {num_factura}: {e}")
        return False

def validar_json_para_envio(json_data, factura_num):
    errores = []

    if "rips" not in json_data:
        errores.append("Falta el campo 'rips'.")
    else:
        rips = json_data["rips"]
        if "numFactura" not in rips:
            errores.append("Falta 'numFactura'.")
        if "usuarios" not in rips or not rips["usuarios"]:
            errores.append("'usuarios' vacío o inexistente.")

    if not json_data.get("xmlFevFile"):
        errores.append("Falta 'xmlFevFile' o está vacío.")

    if errores:
        print(f"❌ Errores en el JSON de la factura {factura_num}:")
        for error in errores:
            print(f"   - {error}")
        return False

    print(f"✅ JSON de factura {factura_num} está listo para enviar.")
    return True

def verificar_xml_base64_para_todas_las_facturas():
    corregidos = [f for f in os.listdir(UPLOAD_FOLDER) if f.endswith("_CORREGIDO.json")]

    for json_file in corregidos:
        json_path = os.path.join(UPLOAD_FOLDER, json_file)

        try:
            with open(json_path, "r", encoding="utf-8") as f:
                data = json.load(f)

            num_factura = data.get("rips", {}).get("numFactura", "Sin número")
            xml_base64 = data.get("xmlFevFile", "")

            if not xml_base64:
                print(f"⚠️ Factura {num_factura}: No tiene XML codificado en Base64.")
                continue

            xml_decoded = base64.b64decode(xml_base64).decode("utf-8")
            print(f"\n🔍 XML decodificado para factura {num_factura}:")
            print(xml_decoded)
            print("-" * 80)

        except Exception as e:
            print(f"❌ Error procesando {json_file}: {e}")

def cargar_historial():
    """
    Carga el historial de corregidas.json, filtra entradas cuyo
    archivo corregido ya no existe, y vuelve a guardar si hubo cambios.
    """
    try:
        with open(HISTORIAL_PATH, 'r', encoding='utf-8') as f:
            historial = json.load(f)
    except Exception:
        historial = {}

    # historial es dict: factura_num → {fecha, observacion}
    entradas_originales = set(historial.keys())
    entradas_validas = {}

    for num, info in historial.items():
        nombre_archivo = f"{num}_2_CUV_CORREGIDO.json"
        ruta_archivo  = os.path.join(UPLOAD_FOLDER, nombre_archivo)
        if os.path.exists(ruta_archivo):
            entradas_validas[num] = info
        else:
            print(f"⚠️ Eliminando factura {num} del historial (archivo {nombre_archivo} no encontrado)")

    # Si hubo eliminación, reescribimos el JSON para mantenerlo en disco sincronizado
    if set(entradas_validas.keys()) != entradas_originales:
        with open(HISTORIAL_PATH, 'w', encoding='utf-8') as f:
            json.dump(entradas_validas, f, indent=2, ensure_ascii=False)

    return entradas_validas

def guardar_historial(historial):
    with open(HISTORIAL_PATH, 'w', encoding='utf-8') as f:
        json.dump(historial, f, indent=2, ensure_ascii=False)

historial_corregidas = cargar_historial()

def enviar_jsons_corregidos():
    historial_corregidas = cargar_historial()
    corregidos = [f for f in os.listdir(UPLOAD_FOLDER) if f.endswith('_CORREGIDO.json')]
    print(f"🔎 Facturas corregidas encontradas: {corregidos}")

    for json_file in corregidos:
        num_factura = json_file.split("_")[0]
        if num_factura in historial_corregidas:
            print(f"⏭️ Factura {num_factura} ya enviada, saltando...")
            continue

        json_path = os.path.join(UPLOAD_FOLDER, json_file)
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            if not validar_json_para_envio(data, data.get('rips', {}).get('numFactura', 'Sin número')):
                continue

            print(f"📡 Enviando {json_file} al Ministerio...")
            response = requests.post(API_CARGA_JSON_URL, json=data, verify=False)

            if response.status_code == 200:
                respuesta = response.json()
                if respuesta.get('ResultState'):
                    print(f"✅ {json_file}: CUV generado correctamente: {respuesta.get('CodigoUnicoValidacion')}")
                    # Agregar al historial la factura ya enviada
                    historial_corregidas.append(num_factura)
                    guardar_historial(historial_corregidas)
                else:
                    motivo = respuesta.get('ResultadosValidacion', [{}])[0].get('Descripcion', 'Desconocido')
                    print(f"❌ {json_file}: Ministerio rechazó. Motivo: {motivo}")
            else:
                print(f"❌ {json_file}: Error HTTP {response.status_code}")

        except Exception as e:
            print(f"⚠️ Error procesando {json_file}: {e}")

# ========== Inicializar Flask App ==========

app = Flask(__name__)
app.secret_key = 'supersecretkey'

# Configura Flask-Session para que guarde la session en el servidor (filesystem)
app.config['SESSION_TYPE'] = 'filesystem'
app.config['SESSION_FILE_DIR'] = os.path.join(os.getcwd(), 'flask_session')
os.makedirs(app.config['SESSION_FILE_DIR'], exist_ok=True)

# Inicializa Flask-Session
Session(app)
       
@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        limpiar_archivos_sin_cuv(UPLOAD_FOLDER)

        TEMP_UPLOADS_FOLDER = r"C:\Users\USUARIO HILA\Documents\sistema_facturas\temp_uploads"
        os.makedirs(TEMP_UPLOADS_FOLDER, exist_ok=True)
        for nombre in os.listdir(TEMP_UPLOADS_FOLDER):
            path = os.path.join(TEMP_UPLOADS_FOLDER, nombre)
            if os.path.isfile(path):
                os.remove(path)
            else:
                shutil.rmtree(path)

        excel = request.files.get("excel")
        carpeta_archivos = request.files.getlist("carpeta")
        if not excel or not carpeta_archivos:
            return render_template("index.html",
                                   mensaje="⚠️ Debes subir el archivo Excel y la carpeta de archivos.")

        excel_path = os.path.join(TEMP_UPLOADS_FOLDER, excel.filename)
        excel.save(excel_path)
        session['excel_path'] = excel_path
        session.modified = True

        try:
            libro = pd.read_excel(excel_path, sheet_name=None)
            facturas_excel = []
            for nombre_hoja, df in libro.items():
                if df.shape[1] < 1:
                    continue
                series = df.iloc[:, 0].dropna().astype(str).str.strip()
                facturas_excel.extend(series.tolist())
            columna_detectada = ", ".join(libro.keys())
        except Exception as e:
            return render_template("index.html",
                                   mensaje=f"Error al leer el Excel: {e}")

        archivo_dict = {}
        for archivo in carpeta_archivos:
            nombre = os.path.basename(archivo.filename)
            ruta = os.path.join(TEMP_UPLOADS_FOLDER, nombre)
            archivo.save(ruta)
            archivo_dict[nombre] = ruta

        try:
            historial = cargar_historial()
        except Exception:
            historial = {}

        facturas_con_error_xml = []
        facturas_con_cuv_valido   = []
        facturas_con_otros_errores = []

        todas = set(facturas_excel)
        for nombre in archivo_dict:
            if nombre.endswith("_2.json") and "CUV_CORREGIDO" not in nombre:
                todas.add(nombre.split("_")[0])
            if nombre.endswith("_2_Error.json"):
                todas.add(nombre.split("_")[0])

        for num in todas:
            if num in historial:
                facturas_con_cuv_valido.append({
                    "factura":     num,
                    "descripcion": "CUV generado correctamente (registro previo)",
                    "observacion": historial[num]["observacion"]
                })
                continue

            ruta_error = archivo_dict.get(f"{num}_2_Error.json")
            if ruta_error and os.path.exists(ruta_error):
                try:
                    with open(ruta_error, encoding="utf-8") as f:
                        data = json.load(f)
                    rv = data.get("ResultadosValidacion", [])

                    if data.get("ResultState") is True:
                        facturas_con_cuv_valido.append({
                            "factura": num,
                            "descripcion": "CUV validado",
                            "observacion": "CUV validado"
                        })
                        continue  

                except Exception as e:
                    facturas_con_otros_errores.append({
                        "factura": num,
                        "descripcion": "Error leyendo JSON de error",
                        "observacion": str(e)
                    })
                    continue

                if rv:
                    cfr = next((r for r in rv if r.get("Codigo")=="CFR006"), None)
                    if cfr:
                        facturas_con_error_xml.append({
                            "factura":     num,
                            "descripcion": cfr.get("Descripcion",""),
                            "observacion": cfr.get("Observaciones","")
                        })
                    elif any(r.get("Clase")=="RECHAZADO" and "[AttachedDocument]" in r.get("Descripcion","") for r in rv):
                        detail = next(r for r in rv if "[AttachedDocument]" in r.get("Descripcion",""))
                        facturas_con_error_xml.append({
                            "factura":     num,
                            "descripcion": detail.get("Descripcion",""),
                            "observacion": detail.get("Observaciones","")
                        })
                    else:
                        first = rv[0]
                        facturas_con_otros_errores.append({
                            "factura":     num,
                            "descripcion": first.get("Descripcion",""),
                            "observacion": first.get("Observaciones","")
                        })
                else:
                    facturas_con_otros_errores.append({
                        "factura":     num,
                        "descripcion": "Sin validaciones",
                        "observacion": "No hay ResultadosValidacion"
                    })
            else:
                ruta_norm = archivo_dict.get(f"{num}_2.json")
                if ruta_norm and os.path.exists(ruta_norm):
                    facturas_con_cuv_valido.append({
                        "factura":     num,
                        "descripcion": "JSON válido (pendiente de corrección)",
                        "observacion": "Pendiente"
                    })
                else:
                    facturas_con_otros_errores.append({
                        "factura":     num,
                        "descripcion": "Sin JSON",
                        "observacion": "No se encontró archivo"
                    })

        facturas_con_error_xml = [
            f for f in facturas_con_error_xml if f["factura"] not in historial
        ]

        session["columna_detectada"]          = columna_detectada
        session["facturas_con_error"]         = facturas_con_error_xml
        session["facturas_con_cuv_corregido"] = facturas_con_cuv_valido
        session["facturas_con_otros_errores"] = facturas_con_otros_errores
        session["archivos_guardados"]         = archivo_dict
        session.modified = True

        return redirect(url_for("resultados"))

    return render_template("index.html")

@app.route("/resultado")
def resultados():
    historial = cargar_historial()
    errores = session.get("facturas_con_error", [])
    cuv_previos = session.get("facturas_con_cuv_corregido", [])
    otros = session.get("facturas_con_otros_errores", [])

    facturas_con_error = [f for f in errores if f["factura"] not in historial]

    facturas_con_cuv_corregido = []
    ya_agregadas = set()

    for f in cuv_previos:
        facturas_con_cuv_corregido.append({
            "factura": f["factura"],
            "descripcion": f.get("descripcion", "CUV generado correctamente"),
            "observacion": "CUV validado"
        })
        ya_agregadas.add(f["factura"])

    for num, info in historial.items():
        if num not in ya_agregadas:
            facturas_con_cuv_corregido.append({
                "factura": num,
                "descripcion": "CUV generado correctamente (registro previo)",
                "observacion": "CUV validado"
            })

    return render_template("resultados.html",
                           columna_detectada=session.get("columna_detectada", ""),
                           facturas_con_error=facturas_con_error,
                           facturas_con_cuv_corregido=facturas_con_cuv_corregido,
                           facturas_con_otros_errores=otros)


@app.route("/vista_excel") 
def vista_excel():
    historial = cargar_historial()
    corregidas = set(historial.keys())
    facturas_set = set()  # Para evitar duplicados
    facturas = []

    def clasificar_descripcion(desc):
        desc = desc.lower()
        if "cuota moderadora" in desc or "pagos moderadores" in desc:
            return "Cuota moderada"
        elif "valor reportado en los servicios" in desc:
            return "Cirugías"
        elif "xml" in desc:
            return "XML"
        else:
            return desc  

    # Facturas con error
    for f in session.get("facturas_con_error", []):
        factura = f["factura"]
        if factura in facturas_set:
            continue
        facturas_set.add(factura)
        estado = "Corregida" if factura in corregidas else "No corregida"
        if estado == "Corregida":
            descripcion = "Factura válida por el Ministerio"
        else:
            descripcion = clasificar_descripcion(f["descripcion"])
        facturas.append({"factura": factura, "estado": estado, "descripcion": descripcion})

    # Facturas válidas
    for f in session.get("facturas_con_cuv_corregido", []):
        factura = f["factura"]
        if factura in facturas_set:
            continue
        facturas_set.add(factura)
        facturas.append({
            "factura": factura,
            "estado": "Válida",
            "descripcion": "Factura válida por el Ministerio"
        })

    # Otros errores
    for f in session.get("facturas_con_otros_errores", []):
        factura = f["factura"]
        if factura in facturas_set:
            continue
        facturas_set.add(factura)
        descripcion = clasificar_descripcion(f["descripcion"])
        facturas.append({
            "factura": factura,
            "estado": "Inválida",
            "descripcion": descripcion
        })

    return render_template("vista_excel.html", facturas=facturas)

@app.route("/descargar_excel_actualizado", methods=["POST"])
def descargar_excel_actualizado():
    import pandas as pd
    import io
    import os
    from flask import session, send_file

    # 1) Recuperar ruta del Excel original de la sesión
    excel_path = session.get("excel_path")
    if not excel_path or not os.path.exists(excel_path):
        return "⚠️ No encuentro el Excel original en sesión.", 400

    # 2) Volver a cargar todas las hojas
    try:
        libro = pd.read_excel(excel_path, sheet_name=None)
    except Exception as e:
        return f"❌ Error al reabrir el Excel: {e}", 500

    # 3) Diccionarios desde sesión
    errores_dict = {f["factura"]: f["descripcion"] for f in session.get("facturas_con_error", [])}
    valido_dict = {f["factura"]: f["observacion"] for f in session.get("facturas_con_cuv_corregido", [])}
    otros_dict = {f["factura"]: f["descripcion"] for f in session.get("facturas_con_otros_errores", [])}

    # 4) Clasificador de descripción
    def clasificar_descripcion(msg):
        msg_lower = msg.lower()
        if "cuotas moderadoras" in msg_lower or "pagos moderadores" in msg_lower:
            return "Cuota moderada"
        elif "valor reportado en los servicios" in msg_lower:
            return "Cirugías"
        elif "xml" in msg_lower or "attacheddocument" in msg_lower:
            return "XML"
        else:
            return msg

    # 5) Generar nuevo Excel con colores
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        for hoja, df_hoja in libro.items():
            if df_hoja.shape[1] < 1:
                continue

            facs = df_hoja.iloc[:, 0].dropna().astype(str).str.strip()
            filas = []

            for fac in facs:
                if fac in errores_dict:
                    estado = "Error"
                    descripcion = clasificar_descripcion(errores_dict[fac])
                elif fac in valido_dict:
                    estado = "Válida"
                    descripcion = "Factura válida por el Ministerio"
                elif fac in otros_dict:
                    estado = "Otro error"
                    descripcion = clasificar_descripcion(otros_dict[fac])
                else:
                    estado = "No procesada"
                    descripcion = ""

                filas.append({
                    "Factura": fac,
                    "Estado": estado,
                    "Descripción": descripcion
                })

            df_resultado = pd.DataFrame(filas)
            df_resultado.to_excel(writer, sheet_name=hoja, index=False)

            workbook  = writer.book
            worksheet = writer.sheets[hoja]

            formato_verde = workbook.add_format({'bg_color': "#30D651"})

            for fila_idx, estado in enumerate(df_resultado["Estado"], start=1):  
                if estado == "Válida":
                    worksheet.set_row(fila_idx, None, formato_verde)

    output.seek(0)
    return send_file(
        output,
        download_name="facturas_actualizadas_por_hoja.xlsx",
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

@app.route('/descargar_factura/<factura_id>')
def descargar_factura(factura_id):
    archivos = []
    for folder in (INCOMING_FOLDER, UPLOAD_FOLDER):
        if not os.path.isdir(folder):
            continue
        for f in os.listdir(folder):
            if f.startswith(f"{factura_id}_"):
                archivos.append((os.path.join(folder, f), f))

    if not archivos:
        return redirect(url_for('vista_excel', error_factura=factura_id))

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zf:
        for ruta_completa, nombre in archivos:
            zf.write(ruta_completa, arcname=nombre)
    zip_buffer.seek(0)

    # 3) Enviar el ZIP 
    return send_file(
        zip_buffer,
        as_attachment=True,
        download_name=f'factura_{factura_id}_archivos.zip',
        mimetype='application/zip'
    )

@app.route('/descargar_todas_facturas', methods=['POST'])
def descargar_todas_facturas():
    facturas = request.form.getlist('facturas[]')  
    archivos = []

    for factura_id in facturas:
        for folder in (INCOMING_FOLDER, UPLOAD_FOLDER):
            if not os.path.isdir(folder):
                continue
            for f in os.listdir(folder):
                if f.startswith(f"{factura_id}_"):
                    archivos.append((os.path.join(folder, f), f))

    if not archivos:
        return "No se encontraron archivos para las facturas seleccionadas", 404

    # Crear el ZIP en memoria
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zf:
        for ruta_completa, nombre in archivos:
            zf.write(ruta_completa, arcname=nombre)
    zip_buffer.seek(0)

    return send_file(
        zip_buffer,
        as_attachment=True,
        download_name='facturas_seleccionadas.zip',
        mimetype='application/zip'
    )

@app.route("/ver_reportes")
def ver_reportes():
    from datetime import datetime

    # 1) Carga el historial corregidas.json
    if os.path.exists(HISTORIAL_PATH):
        with open(HISTORIAL_PATH, "r", encoding="utf-8") as f:
            historial_corregidas = json.load(f)
    else:
        historial_corregidas = {}

    # 2) Agrupa las facturas por mes_clave = "YYYY-MM"
    facturas_por_mes = defaultdict(list)
    for factura_id, info in historial_corregidas.items():
        fecha = info.get("fecha", "0000-00-00")
        mes_clave = fecha[:7]  
        facturas_por_mes[mes_clave].append({
            "factura": factura_id,
            "observacion": info.get("observacion", ""),
            "fecha": fecha
        })

    # 3) Extrae años y meses únicos para los filtros
    anos = sorted({m[:4] for m in facturas_por_mes.keys()})
    meses = sorted({m[5:7] for m in facturas_por_mes.keys()})

    return render_template(
        "reportes.html",
        fecha_actual = datetime.now().strftime("%Y-%m-%d %H:%M"),
        reportes_por_mes = facturas_por_mes,
        anos = anos,
        meses = meses
    )

@app.route("/descargar_pdf/<mes>")
def descargar_pdf_mes(mes):
    if os.path.exists(HISTORIAL_PATH):
        with open(HISTORIAL_PATH, "r", encoding="utf-8") as f:
            historial = json.load(f)
    else:
        historial = {}

    facturas_mes = []
    for num, info in historial.items():
        fecha = info.get("fecha", "")
        if fecha.startswith(mes):
            facturas_mes.append((num, info))

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=40, rightMargin=40, topMargin=80, bottomMargin=40)
    styles = getSampleStyleSheet()
    elements = []

    # === DATOS EMPRESA ===
    elements.append(Paragraph(
        "<b>Hospital Infantil Los Ángeles</b><br/>"
        "NIT: 891.200.240-2<br/>"
        "Cra. 32, Pasto, Nariño<br/>"
        "subgestioninformacion@hinfantil.org",
        styles['Normal']
    ))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(f"📄 <b>Reporte de Facturas Corregidas - {mes}</b>", styles['Title']))
    elements.append(Spacer(1, 6))
    elements.append(Paragraph(
        f"<b>Fecha de generación:</b> {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        styles['Normal']
    ))
    elements.append(Spacer(1, 12))

    if not facturas_mes:
        elements.append(Paragraph("❌ No hay facturas corregidas para este mes.", styles['Normal']))
    else:
        data = [["Factura", "Fecha", "Observación"]]
        for num, info in facturas_mes:
            data.append([str(num), info.get("fecha", ""), "Error XML corregido"])

        table = Table(data, colWidths=[100, 150, 250])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#007bff")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey]),
        ]))
        elements.append(table)

    logo_path = os.path.join("static", "img", "logo.jpg")

    def draw_header(canvas_obj, doc_obj):
        if os.path.exists(logo_path):
            w = 1.5 * inch
            h = (50 / 120) * w
            x = doc_obj.pagesize[0] - doc_obj.rightMargin - w
            y = doc_obj.pagesize[1] - h - 10
            canvas_obj.drawImage(logo_path, x, y, width=w, height=h, preserveAspectRatio=True)
        footer_text = "Sistema de Corrección de XML — Hospital Infantil Los Ángeles"
        canvas_obj.setFont("Helvetica-Oblique", 9)
        canvas_obj.drawCentredString(doc_obj.pagesize[0] / 2.0, 20, footer_text)

    doc.build(elements, onFirstPage=draw_header, onLaterPages=draw_header)
    buffer.seek(0)
    return send_file(
        buffer,
        download_name=f"reporte_facturas_{mes}.pdf",
        as_attachment=True,
        mimetype="application/pdf"
    )

@app.route("/descargar_word/<mes>")
def descargar_word_mes(mes):
    if os.path.exists(HISTORIAL_PATH):
        with open(HISTORIAL_PATH, "r", encoding="utf-8") as f:
            historial = json.load(f)
    else:
        historial = {}

    facturas_mes = []
    for num, info in historial.items():
        fecha = info.get("fecha", "")
        if fecha.startswith(mes):
            facturas_mes.append((num, info))

    doc = Document()
    section = doc.sections[0]

    # === LOGO ===
    header = section.header
    logo_path = os.path.join("static", "img", "logo.jpg")
    if os.path.exists(logo_path):
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(1.0))
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # === DATOS EMPRESA ===
    datos = doc.add_paragraph()
    datos.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = datos.add_run(
        "Hospital Infantil Los Ángeles\n"
        "NIT: 891.200.240-2\n"
        "Cra. 32, Pasto, Nariño\n"
        "subgestioninformacion@hinfantil.org"
    )
    run.font.size = Pt(10)
    run.font.italic = True

    # === TÍTULO ===
    titulo = doc.add_heading(f"📄 Reporte de Facturas Corregidas - {mes}", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.runs[0]
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.size = Pt(20)

    # === FECHA ===
    fecha_p = doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    if not facturas_mes:
        doc.add_paragraph("❌ No hay facturas corregidas para este mes.")
    else:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Factura"
        hdr_cells[1].text = "Fecha"
        hdr_cells[2].text = "Observación"

        for cell in hdr_cells:
            for p in cell.paragraphs:
                p.runs[0].bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i, (num, info) in enumerate(facturas_mes):
            row_cells = table.add_row().cells
            row_cells[0].text = str(num)
            row_cells[1].text = info.get("fecha", "")
            row_cells[2].text = "Error XML corregido"

            for cell in row_cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if i % 2 == 0:
                for cell in row_cells:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), "D9E1F2")
                    cell._tc.get_or_add_tcPr().append(shading_elm)

    footer = section.footer
    p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p_footer.text = "Sistema de Corrección de XML — Hospital Infantil Los Ángeles"
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_footer.runs[0].italic = True
    p_footer.runs[0].font.size = Pt(9)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return send_file(
        output,
        download_name=f"reporte_facturas_{mes}.docx",
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route('/manual')
def ver_manual():
    return render_template('manual.html')

# 2) Ruta para descargar el PDF del manual
@app.route("/descargar_manual")
def descargar_manual():
    # Construye la ruta al PDF en static/manual
    manual_path = os.path.join(app.root_path, "static", "manual", "manual_usuario.pdf")
    if not os.path.exists(manual_path):
        return "❌ Manual no encontrado", 404

    return send_file(
        manual_path,
        as_attachment=True,
        download_name="manual_usuario.pdf",
        mimetype="application/pdf"
    )

@app.route("/soporte")
def soporte():
    return render_template("soporte.html")

# Ruta para procesar las facturas con error en XML
@app.route("/corregir", methods=["POST"])
def corregir_y_enviar():
    global historial_corregidas
    try:
        print("=== INICIO CORRECCIÓN AUTOMÁTICA ===")

        # 1) Filtrar sólo facturas aún no corregidas
        historial_corregidas = cargar_historial()  # dict: factura → {fecha, observacion}
        facturas_con_error = session.get("facturas_con_error", [])
        facturas_con_error = [
            f for f in facturas_con_error
            if f["factura"] not in historial_corregidas
        ]

        if not facturas_con_error:
            print("⚠️ No hay facturas con error XML nuevas para corregir.")
            return jsonify({"mensaje": "No hay facturas con error XML."}), 400

        archivo_dict = session.get("archivos_guardados", {})
        facturas_cuv = session.get("facturas_con_cuv_corregido", [])

        # 2) Autenticación en API
        r_login = requests.post(API_LOGIN_URL, json=API_CREDENCIALES, verify=False)
        r_login.raise_for_status()
        token = r_login.json().get("token")
        if not token:
            return jsonify({"mensaje": "No se pudo autenticar con el Ministerio."}), 401
        headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}

        errores = []
        corregidas = []

        for f in facturas_con_error:
            num = f["factura"]
            print(f"\n--- Procesando factura {num} ---")

            # 3) Ruta del JSON original
            path_json = archivo_dict.get(f"{num}_2.json")
            if not path_json or not os.path.exists(path_json):
                errores.append(f"{num}: JSON original no encontrado.")
                continue

            # 4) Crear JSON corregido en temp (INCOMING_FOLDER)
            salida = os.path.join(INCOMING_FOLDER, f"{num}_2_CORREGIDO.json")
            num_limpio = limpiar_num_factura(num)
            if not corregir_json_valido(path_json, salida, INCOMING_FOLDER, num_limpio):
                errores.append(f"{num}: Error al corregir el JSON.")
                continue

            # 5) Copiar TODOS los archivos <num>_* de INCOMING_FOLDER → UPLOAD_FOLDER
            import shutil
            for nombre in os.listdir(INCOMING_FOLDER):
                if nombre.startswith(f"{num}_"):
                    src = os.path.join(INCOMING_FOLDER, nombre)
                    dst = os.path.join(UPLOAD_FOLDER, nombre)
                    shutil.copy2(src, dst)
                    print(f"📂 Copiado a uploads/: {nombre}")

            # 6) Validar que el JSON corregido está listo
            with open(salida, "r", encoding="utf-8") as fcor:
                json_corregido = json.load(fcor)
            if not validar_json_para_envio(json_corregido, num):
                errores.append(f"{num}: JSON inválido, no se envió.")
                continue

            # 7) Enviar al Ministerio
            r = requests.post(API_CARGA_JSON_URL, headers=headers,
                              json=json_corregido, verify=False)
            res = r.json()

            # 8) Manejo de CUV existente (RVG02)
            if (not res.get("ResultState")
                and any(i.get("Codigo")=="RVG02" for i in res.get("ResultadosValidacion", []))):
                texto = next(i["Observaciones"]
                             for i in res["ResultadosValidacion"]
                             if i.get("Codigo")=="RVG02")
                m = re.search(r"CUV\s*([0-9a-f]+)", texto)
                if m:
                    cuv = m.group(1)
                    # Armamos un resultado válido
                    nuevo_res = {
                        **{k:v for k,v in res.items() if k!="ResultadosValidacion"},
                        "CodigoUnicoValidacion": cuv,
                        "ResultadosValidacion": [
                            i for i in res["ResultadosValidacion"]
                            if i.get("Clase")=="NOTIFICACION"
                        ],
                        "ResultState": True
                    }
                    res = nuevo_res

            # 9) Si quedó validado (nuevo o existente)
            if res.get("ResultState"):
                # Guardar JSON final en uploads/
                with open(os.path.join(UPLOAD_FOLDER, f"{num}_2_CUV_CORREGIDO.json"),
                          "w", encoding="utf-8") as f_cuv:
                    json.dump(res, f_cuv, indent=2, ensure_ascii=False)
                historial_corregidas[num] = {
                    "fecha": datetime.now().strftime("%Y-%m-%d %H:%M"),
                    "observacion": res.get("CodigoUnicoValidacion","")
                }
                corregidas.append({"factura": num, "observacion": "CUV validado"})
                facturas_cuv.append({
                    "factura": num,
                    "descripcion": "CUV generado correctamente",
                    "observacion": "CUV validado"
                })
                print(f"✅ CUV generado y guardado para {num}")
            else:
                obs = res.get("ResultadosValidacion", [{}])[0].get("Observaciones", "Error desconocido")
                errores.append(f"{num}: {obs}")
                print(f"❌ Ministerio rechazó {num}: {obs}")

        # 10) Actualizar sesión e historial, limpiar obsoletos
        facturas_con_error = [
            f for f in facturas_con_error
            if f["factura"] not in [c["factura"] for c in corregidas]
        ]
        session["facturas_con_error"] = facturas_con_error
        session["facturas_con_cuv_corregido"] = facturas_cuv
        session.modified = True

        guardar_historial(historial_corregidas)
        limpiar_archivos_sin_cuv(UPLOAD_FOLDER)

        return jsonify({
            "mensaje": "Corrección finalizada",
            "corregidas": corregidas,
            "errores": errores,
            "total_corregidas": len(corregidas),
            "total_no_corregidas": len(errores)
        }), 200

    except Exception as e:
        print(f"❌ ERROR GENERAL: {e}")
        return jsonify({"mensaje": "Error inesperado"}), 500

if __name__ == "__main__":
    app.run(debug=True)